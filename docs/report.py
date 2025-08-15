import os
from docx import Document

# Укажите путь к вашей папке
folder_path = '/home/maxat/Projects/Agrarka/scales-installer/submodule/tests'

# Создаем новый документ
report = Document()
report.add_heading('Объединённый отчет', 0)

# Проходим по всем файлам в папке
for filename in sorted(os.listdir(folder_path)):
    file_path = os.path.join(folder_path, filename)
    if os.path.isfile(file_path):
        report.add_heading(filename, level=1)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                report.add_paragraph(content)
        except Exception as e:
            report.add_paragraph(f"Ошибка при чтении {filename}: {e}")

# Сохраняем итоговый файл
report.save('tests.docx')
